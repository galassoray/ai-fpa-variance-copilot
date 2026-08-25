"""
test_agent_reports.py
=====================
The Word deliverables, and the rule they inherit from the deck.

THE PROPERTY
------------
No number can appear in a document unless a tool returned it.

The test re-opens each generated ``.docx``, reads every paragraph and every
table cell back out, and audits them with the SAME numeric auditor used on the
commentary. Checking the generator against its own provenance list would prove
only that it is self-consistent.

This matters here more than anywhere else in the project. Six comparative tools
were added and their fields went unregistered downstream; nothing failed, and a
month-over-month question ran perfectly and produced an empty answer. A
document generator is a new downstream surface with exactly the same exposure,
so it gets the same mechanism rather than the same care.
"""
import os
import re
import sys

import pytest

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "..", "src")
sys.path.insert(0, SRC)

pytest.importorskip("docx")

from docx import Document  # noqa: E402

import run_pipeline as rp  # noqa: E402
from agent import materialize as mz  # noqa: E402
from agent import tools as _tools  # noqa: E402,F401
from agent.facts import fact_pack_from_ledger  # noqa: E402
from agent.gates import PublicationPacket  # noqa: E402
from agent.narrate import narrate  # noqa: E402
from agent.orchestrator import Orchestrator  # noqa: E402
from agent.packages import build_goal, variance_package_plan  # noqa: E402
from agent.plan import Plan, Step  # noqa: E402
from agent.reports import (BLANK_NOTE, Report, ReportError, build_all,  # noqa: E402
                           build_flash, build_memo, build_packet,
                           department_ids, fmt)
from guardrails import entity_audit as ea  # noqa: E402
from guardrails import numeric_audit as na  # noqa: E402

PERIOD = "2025-09"


@pytest.fixture(scope="module")
def con():
    mz.materialize(rp.compute(rp.load()), verbose=False)
    c = mz.connect_readonly()
    yield c
    c.close()


@pytest.fixture(scope="module")
def goal(con):
    return build_goal(con, PERIOD)


@pytest.fixture(scope="module")
def run(con, goal):
    return Orchestrator(con).run(variance_package_plan(goal), goal)


@pytest.fixture(scope="module")
def candidate(run, goal):
    return narrate(run, goal, None, ea.canonical_entity_names(rp.load()),
                   mode="inject")


@pytest.fixture(scope="module")
def packet(run, candidate):
    p = PublicationPacket(run, candidate)
    p.approve("Test Reviewer", "acceptance run")
    return p


@pytest.fixture(scope="module")
def written(tmp_path_factory, run, goal, candidate, packet):
    out = tmp_path_factory.mktemp("reports")
    return build_all(run, goal, candidate, packet, str(out))


def doc_text(path: str) -> str:
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


#: Counts and ordinals that are document furniture rather than figures:
#: driver ranks, the verified-figure count, years inside a date.
STRUCTURAL = set(range(0, 30)) | set(range(1900, 2101))


def structural(value: float) -> bool:
    return float(value).is_integer() and int(value) in STRUCTURAL


# --------------------------------------------------------------------------
# 1. the property
# --------------------------------------------------------------------------
def test_every_document_is_produced(written):
    names = sorted(os.path.basename(p) for p in written)
    assert any(n.startswith("flash-") for n in names)
    assert any(n.startswith("memo-") for n in names)
    packets = [n for n in names if n.startswith("packet-")]
    assert len(packets) == 5, f"expected one packet per department: {names}"


@pytest.mark.parametrize("kind", ["flash", "memo", "packet"])
def test_every_number_in_every_document_traces_to_the_ledger(
        kind, written, run, goal):
    """Reads the generated FILE and audits it with the same auditor used on
    the commentary."""
    pack = fact_pack_from_ledger(run, goal)
    paths = [p for p in written if os.path.basename(p).startswith(kind)]
    assert paths

    failures = []
    for path in paths:
        result = na.audit(doc_text(path), pack)
        for v in result.violations:
            if structural(v.value):
                continue
            failures.append((os.path.basename(path), v.mention, v.reason))

    assert not failures, (
        "figures with no computed source: "
        + "; ".join(f"{f}: {m} ({r})" for f, m, r in failures[:8])
    )


def test_the_generator_performs_no_arithmetic_on_figures():
    """Enforced by reading the source, so it survives future edits.

    A formatter that sums a column or derives a percentage becomes a second
    path to a number -- the failure every layer beneath this one exists to
    prevent, arriving in the most circulated artifact.
    """
    src = open(os.path.join(SRC, "agent", "reports.py"), encoding="utf-8").read()
    marker = "# FLASH"
    assert marker in src
    body = "\n".join(ln for ln in src[src.index(marker):].splitlines()
                     if not ln.strip().startswith("#"))
    banned = (r"\b(sum|mean|round|abs|min|max)\s*\(|\bstatistics\.|\bnp\.|"
              r"\bmath\.")
    hits = [ln.strip() for ln in body.splitlines() if re.search(banned, ln)]
    assert not hits, f"the document layer derives a figure: {hits[:3]}"


def test_provenance_records_one_entry_per_figure(run, goal, candidate, packet):
    r = build_memo(run, goal, candidate, packet)
    assert r.provenance
    for ref in r.provenance:
        assert ref.section in run.sections
        assert ref.step > 0 and ref.display


def test_fig_refuses_a_value_that_was_not_retrieved(run, goal):
    r = Report(run, goal)
    with pytest.raises(ReportError, match="not in this run"):
        r.fig("no_such_section", 0, "actual")
    section = sorted(run.sections)[0]
    with pytest.raises(ReportError, match="row 99 was requested"):
        r.fig(section, 99, "actual")
    with pytest.raises(ReportError, match="not a field"):
        r.fig(section, 0, "invented_field")


# --------------------------------------------------------------------------
# 2. the blank recommendation
# --------------------------------------------------------------------------
@pytest.mark.parametrize("kind", ["flash", "memo", "packet"])
def test_every_document_leaves_the_recommendation_to_the_analyst(kind, written):
    """The standard variance-commentary format ends with a recommendation.
    This tool does not generate one -- a recommendation is not a retrieved
    number, so nothing could verify it. The section is present and explicitly
    empty, which makes the division of labour visible on the page."""
    for path in [p for p in written if os.path.basename(p).startswith(kind)]:
        text = doc_text(path)
        assert "Analyst assessment and recommended action" in text
        assert BLANK_NOTE in text
        assert "judgment belongs to the analyst" in text


@pytest.mark.parametrize("kind", ["flash", "memo", "packet"])
def test_no_document_states_a_recommendation(kind, written):
    """The prose must not quietly advise while the box claims it does not."""
    advice = ["we recommend", "you should", "consider cutting",
              "should reduce", "should increase", "must cut"]
    for path in [p for p in written if os.path.basename(p).startswith(kind)]:
        low = doc_text(path).lower()
        for phrase in advice:
            assert phrase not in low, f"{os.path.basename(path)}: {phrase!r}"


# --------------------------------------------------------------------------
# 3. the packet
# --------------------------------------------------------------------------
def test_a_packet_contains_only_its_own_department(written, run):
    """The point of a budget-owner packet is that an owner sees their own
    numbers. A packet carrying another department's account detail would be
    both confusing and, in a real organisation, a disclosure problem."""
    names = {}
    for name, sec in run.sections.items():
        if sec["tool"] == "decompose_variance":
            dept = sec["params"]["department_id"]
            names[dept] = {r["account_name"] for r in sec["rows"]}

    for path in [p for p in written if os.path.basename(p).startswith("packet")]:
        dept = os.path.basename(path).rsplit("-", 1)[1].split(".")[0]
        own = names.get(dept, set())

        # Exact cell values, not substring containment. Substring matching
        # reported R&D's "Software" as leaking into the S&M packet, because
        # S&M has "Sales Software" -- a false positive in the test, not a
        # disclosure in the document.
        cells = set()
        for t in Document(path).tables:
            for row in t.rows:
                for cell in row.cells:
                    cells.add(cell.text.strip())

        for other, accounts in names.items():
            if other == dept:
                continue
            leaked = sorted((accounts - own) & cells)
            assert not leaked, (
                f"{os.path.basename(path)} shows {other} accounts: {leaked}"
            )


def test_every_department_gets_a_packet(run, goal, candidate, tmp_path):
    """The plan decomposes ALL departments precisely so that every budget
    owner can be served, not just the two largest drivers."""
    depts = department_ids(run)
    assert len(depts) == 5
    for d in depts:
        r = build_packet(run, goal, d, candidate)
        text_path = r.save(str(tmp_path / f"{d}.docx"))
        text = doc_text(text_path)
        assert "Account detail" in text
        assert "Sign-off" in text


def test_a_packet_carries_a_signature_block(written):
    for path in [p for p in written if os.path.basename(p).startswith("packet")]:
        text = doc_text(path)
        assert "Sign-off" in text
        assert "Budget owner" in text and "Prepared by (FP&A)" in text


def test_a_packet_states_the_operating_income_basis(written):
    """Account detail is on an operating-income basis and the People block is
    on an expense basis. A budget owner reading both must be told which is
    which -- the same sign confusion that produced a wrong claim earlier.

    The People block is asserted only where it EXISTS: a revenue-holding
    department has no salary lines, and demanding the caption there would be
    demanding a label for a table that is correctly absent.
    """
    for path in [p for p in written if os.path.basename(p).startswith("packet")]:
        text = doc_text(path)
        assert "impact on operating income" in text.lower()
        if "People" in text:
            assert "expense basis" in text.lower(), (
                f"{os.path.basename(path)} shows a People block without "
                "stating its basis"
            )


# --------------------------------------------------------------------------
# 4. the documents follow the plan
# --------------------------------------------------------------------------
def test_a_thin_plan_produces_a_shorter_document_not_a_broken_one(
        con, goal, tmp_path):
    """Same property as the deck and the narrative: the plan determines the
    output, and a block whose data was never retrieved is omitted rather than
    left as an empty placeholder."""
    thin = Orchestrator(con).run(Plan(goal="g", steps=[
        Step(1, "get_operating_metrics", {"period": "$GOAL.period"},
             purpose="o"),
        Step(2, "get_pl_summary", {"period": "$GOAL.period"}, purpose="p"),
    ], promises=["o", "p"]), goal)

    path = build_flash(thin, goal).save(str(tmp_path / "thin.docx"))
    text = doc_text(path)
    assert "Headline" in text and "Versus plan" in text
    assert "Compensation" not in text


def test_documents_locate_sections_by_tool_not_by_label(con, goal, tmp_path):
    """Section labels are planner-chosen, so an agent-authored plan must still
    produce a full memo."""
    agent = Orchestrator(con).run(Plan(goal="g", steps=[
        Step(1, "get_pl_summary", {"period": "$GOAL.period"}, purpose="zzz"),
        Step(2, "rank_variance_drivers",
             {"period": "$GOAL.period", "dimension": "department", "top_n": 5},
             purpose="aaa"),
        Step(3, "get_operating_metrics", {"period": "$GOAL.period"},
             purpose="mmm"),
    ], promises=["zzz", "aaa", "mmm"]), goal)

    text = doc_text(build_memo(agent, goal).save(str(tmp_path / "agent.docx")))
    assert "Results versus plan" in text
    assert "Where the variance came from" in text


@pytest.mark.parametrize("period", ["2024-06", "2025-03", "2025-12"])
def test_documents_build_for_every_period(con, period, tmp_path):
    g = build_goal(con, period)
    res = Orchestrator(con).run(variance_package_plan(g), g)
    cand = narrate(res, g, None, ea.canonical_entity_names(rp.load()),
                   mode="inject")
    paths = build_all(res, g, cand, None, str(tmp_path / period))
    assert len(paths) == 7
    for p in paths:
        assert Document(p).paragraphs


# --------------------------------------------------------------------------
# 5. presentation
# --------------------------------------------------------------------------
def test_every_document_states_the_data_is_synthetic(written):
    """A clearance line, not a caption. These are the artifacts that get
    circulated, so it must appear without being sought."""
    for path in written:
        assert "Synthetic data" in doc_text(path)


def test_negatives_use_accounting_parentheses():
    assert fmt(-142610.66) == "($142,611)"
    assert fmt(15559.0) == "$15,559"
    assert fmt(0.452, "percent") == "45.2%"
    assert fmt(4, "count_signed") == "+4"
    assert fmt(None) == "\u2014"


def test_no_table_has_an_unlabelled_header(written):
    """An empty header cell reads as a formatting mistake to the person being
    asked to sign underneath it."""
    for path in written:
        for t in Document(path).tables:
            headers = [c.text.strip() for c in t.rows[0].cells]
            if len(headers) > 1 and any(headers):
                assert headers[0], (
                    f"{os.path.basename(path)}: table with a blank first header"
                )


# --------------------------------------------------------------------------
# 6. CLI and app integration
# --------------------------------------------------------------------------
def test_run_package_reports_flag(tmp_path, capsys, con):
    from agent import run_package as cli

    out = str(tmp_path / "docs")
    assert cli.main([PERIOD, "--reports", out]) == 0
    printed = capsys.readouterr().out
    assert "7 document(s) written" in printed
    assert len(os.listdir(out)) == 7


def test_the_app_offers_the_documents_after_sign_off(monkeypatch):
    """Deliverables are gated on sign-off: an unsigned packet is exactly the
    artifact that should not leave the building."""
    pytest.importorskip("streamlit")
    from streamlit.testing.v1 import AppTest

    for env in ("OPENAI_API_KEY", "ANTHROPIC_API_KEY"):
        monkeypatch.delenv(env, raising=False)
    sys.path.insert(0, os.path.join(HERE, "..", "eval"))

    at = AppTest.from_file(os.path.join(HERE, "..", "app", "app.py"),
                           default_timeout=240)
    at.run()
    at = at.sidebar.radio[0].set_value("Close-cycle agent").run()
    assert not at.exception

    labels = [b.label for b in at.button]
    for expected in ("Build the deck", "Build the flash", "Build the memo",
                     "Build all 5 packets"):
        assert expected not in labels, (
            f"'{expected}' must not be offered before sign-off"
        )

    name = next(t for t in at.text_input if t.label == "Your name")
    at = name.set_value("Test Reviewer").run()
    at = next(b for b in at.button if b.label == "Sign off").click().run()
    assert not at.exception

    # All four deliverables live under one "Deliverables" heading, each
    # explaining what it produces before it is pressed.
    for label in ("Build the deck", "Build the flash", "Build the memo",
                  "Build all 5 packets"):
        assert any(b.label == label for b in at.button), f"missing: {label}"
        at = next(b for b in at.button if b.label == label).click().run()
        assert not at.exception, f"{label}: {at.exception}"

    downloads = [d.label for d in at.get("download_button")]
    assert any(d.endswith(".pptx") for d in downloads)
    assert any(d.startswith("flash-") for d in downloads)
    assert any(d.startswith("memo-") for d in downloads)
    assert sum(1 for d in downloads if d.startswith("packet-")) == 5


def test_the_deliverables_explain_themselves_before_being_pressed(monkeypatch):
    """A reader should know what a button produces without pressing it.

    The deck and the documents previously sat in separate sections, so a
    visitor had to discover that the tool made four different things.
    """
    pytest.importorskip("streamlit")
    from streamlit.testing.v1 import AppTest

    for env in ("OPENAI_API_KEY", "ANTHROPIC_API_KEY"):
        monkeypatch.delenv(env, raising=False)
    sys.path.insert(0, os.path.join(HERE, "..", "eval"))

    at = AppTest.from_file(os.path.join(HERE, "..", "app", "app.py"),
                           default_timeout=240)
    at.run()
    at = at.sidebar.radio[0].set_value("Close-cycle agent").run()
    at = next(t for t in at.text_input
              if t.label == "Your name").set_value("Reviewer").run()
    at = next(b for b in at.button if b.label == "Sign off").click().run()

    text = " ".join(re.sub("<[^>]+>", "", m.value) for m in at.markdown)
    assert "Deliverables" in text
    for heading in ("Board deck (PowerPoint)", "Flash results (Word)",
                    "Monthly variance commentary (Word)",
                    "Budget owner packets"):
        assert heading in text, f"missing heading: {heading}"

    assert "native PowerPoint objects" in text
    assert "only that owner's numbers" in text
    assert "intentionally blank" in text.lower()


def test_sign_off_does_not_leak_between_sessions(monkeypatch):
    """Regression for a genuine multi-user defect on the deployed app.

    The deterministic run is cached with st.cache_resource, which is shared
    across every session on the server. The cache originally returned the
    PublicationPacket too -- so one visitor signing off made the package appear
    signed to everyone else, on the artifact whose entire purpose is that a
    named human accepted it.

    Found by a test that signed off in one app instance and saw the documents
    already unlocked in the next. The run is shared; the decision is not.
    """
    pytest.importorskip("streamlit")
    from streamlit.testing.v1 import AppTest

    for env in ("OPENAI_API_KEY", "ANTHROPIC_API_KEY"):
        monkeypatch.delenv(env, raising=False)
    sys.path.insert(0, os.path.join(HERE, "..", "eval"))
    app_path = os.path.join(HERE, "..", "app", "app.py")

    def open_agent():
        at = AppTest.from_file(app_path, default_timeout=240)
        at.run()
        return at.sidebar.radio[0].set_value("Close-cycle agent").run()

    first = open_agent()
    name = next(t for t in first.text_input if t.label == "Your name")
    first = name.set_value("Reviewer One").run()
    first = next(b for b in first.button if b.label == "Sign off").click().run()
    assert any("Signed off" in s.value for s in first.success)

    second = open_agent()
    assert not any("Signed off" in s.value for s in second.success), (
        "a sign-off in one session must not appear in another"
    )
    assert any(b.label == "Sign off" for b in second.button)
    assert not any(b.label == "Build the Word documents" for b in second.button)


# --------------------------------------------------------------------------
# 7. no figure is ever truncated on screen
# --------------------------------------------------------------------------
def test_metric_rows_size_themselves_to_their_longest_value():
    """The layout rule, checked directly.

    Streamlit truncates a metric value that does not fit its column with an
    ellipsis -- silently, no wrap, no error. The constraint is COLUMNS, not
    characters: a four-column row holds an eleven-character value and a
    five-column row does not.
    """
    sys.path.insert(0, os.path.join(HERE, "..", "app"))
    src = open(os.path.join(HERE, "..", "app", "app.py"),
               encoding="utf-8").read()
    assert "def _metric_row(items):" in src

    ns = {}
    start = src.index("def _metric_row(items):")
    end = src.index("def _render_package_visual")
    placed = []

    class _Col:
        def metric(self, label, value):
            placed.append((label, value))

    def _columns(n):
        placed.append(("__cols__", n))
        return [_Col() for _ in range(n)]

    ns["st"] = type("st", (), {"columns": staticmethod(_columns)})
    exec(src[start:end], ns)
    row = ns["_metric_row"]

    row([("a", "$1"), ("b", "$2"), ("c", "$3"), ("d", "$4")])
    assert ("__cols__", 4) in placed, "short values should use four columns"

    placed.clear()
    row([("a", "$28,501,685"), ("b", "$28,809,278")])
    assert ("__cols__", 2) in placed

    placed.clear()
    row([("a", "$1,234,567,890,123"), ("b", "$2"), ("c", "$3")])
    widths = [n for lab, n in placed if lab == "__cols__"]
    assert widths and max(widths) <= 3, (
        f"a very long value must narrow the row, got {widths}"
    )

    placed.clear()
    row([("a", "$1"), ("b", None), ("c", "$3")])
    labels = [lab for lab, _ in placed if lab != "__cols__"]
    assert labels == ["a", "c"], "absent values must be dropped, not blanked"


def test_no_metric_row_uses_more_than_four_columns():
    """Enforced by reading the source, so it survives future edits."""
    src = open(os.path.join(HERE, "..", "app", "app.py"),
               encoding="utf-8").read()
    marker = "def _render_package_visual"
    body = src[src.index(marker):]
    body = body[:body.index("\ndef ", 10)]
    wide = [int(n) for n in re.findall(r"st\.columns\((\d+)\)", body)]
    assert all(n <= 4 for n in wide), f"metric row too wide: {wide}"


@pytest.mark.parametrize("period", ["2024-01", "2024-06", "2025-09", "2025-12"])
def test_no_figure_is_truncated_on_screen(period, monkeypatch):
    """Rendered across periods, because the risk grows with the numbers: ARR
    climbs across the dataset, so a layout that fits in January may not in
    December."""
    pytest.importorskip("streamlit")
    from streamlit.testing.v1 import AppTest

    for env in ("OPENAI_API_KEY", "ANTHROPIC_API_KEY"):
        monkeypatch.delenv(env, raising=False)
    sys.path.insert(0, os.path.join(HERE, "..", "eval"))

    at = AppTest.from_file(os.path.join(HERE, "..", "app", "app.py"),
                           default_timeout=300)
    at.run()
    at = at.sidebar.radio[0].set_value("Close-cycle agent").run()
    at = at.sidebar.selectbox[0].set_value(f"{period}-01").run()
    assert not at.exception

    for m in at.metric:
        value = str(m.value)
        assert "\u2026" not in value and "..." not in value, (
            f"{period}: {m.label} is truncated -> {value}"
        )
        assert len(value) <= 12, (
            f"{period}: {m.label} = {value} is long enough to truncate "
            "in a four-column row"
        )


def test_the_arr_block_shows_balances_in_full(monkeypatch):
    """Regression for the reported truncation: five ARR metrics in one row cut
    the eight-figure balances to '$28,501,6...'."""
    pytest.importorskip("streamlit")
    from streamlit.testing.v1 import AppTest

    for env in ("OPENAI_API_KEY", "ANTHROPIC_API_KEY"):
        monkeypatch.delenv(env, raising=False)
    sys.path.insert(0, os.path.join(HERE, "..", "eval"))

    at = AppTest.from_file(os.path.join(HERE, "..", "app", "app.py"),
                           default_timeout=300)
    at.run()
    at = at.sidebar.radio[0].set_value("Close-cycle agent").run()

    labels = {m.label: str(m.value) for m in at.metric}
    assert "Starting ARR" in labels and "Ending ARR" in labels
    for flow in ("New", "Expansion", "Churn"):
        assert flow in labels, f"missing ARR flow: {flow}"
    for lab in ("Starting ARR", "Ending ARR"):
        assert labels[lab].count(",") >= 2, (
            f"{lab} does not look like a full figure: {labels[lab]}"
        )


# --------------------------------------------------------------------------
# 8. a fresh deploy has no database
# --------------------------------------------------------------------------
def test_the_agent_builds_its_own_marts_on_a_cold_start(tmp_path, monkeypatch):
    """Regression for a deployment failure that local runs could never catch.

    data/processed/*.duckdb is gitignored -- correctly, it is a build artifact
    derived from the committed CSVs. So a fresh checkout has no database, and
    the agent opens it READ-ONLY, which cannot create one. Every other page
    computes from the CSVs in memory and never touches DuckDB, so the deployed
    app failed on the agent page alone with "database does not exist" while
    everything else worked.

    Nobody who had ever run the project locally would see it: the file is
    already there.
    """
    from agent import materialize as agent_mz

    fresh = tmp_path / "processed" / "fpa.duckdb"
    monkeypatch.setattr(agent_mz, "DB", str(fresh))
    assert not fresh.exists()

    # Read-only cannot create the database -- this is the deployed failure.
    with pytest.raises(Exception):
        agent_mz.connect_readonly()

    stamp = agent_mz.ensure_ready(verbose=False)
    assert stamp and fresh.exists()

    con = agent_mz.connect_readonly()
    try:
        rows = con.execute("SELECT COUNT(*) FROM out_variance_detail").fetchone()[0]
        assert rows > 0
        agent_mz.assert_fresh(con)
    finally:
        con.close()

    # Cheap and safe to call again: one hash comparison, no rebuild.
    assert agent_mz.ensure_ready(verbose=False) == stamp


def test_the_committed_demo_runs_are_not_gitignored():
    """The deployed app serves these with no key. If they were ignored, the
    public link would have nothing to replay."""
    import subprocess

    repo = os.path.join(HERE, "..")
    runs = os.path.join(repo, "data", "agent_runs")
    assert os.path.isdir(runs), "no committed demo runs"
    files = [f for f in os.listdir(runs) if f.endswith(".json")]
    assert files, "the demo run directory is empty"

    for f in files:
        r = subprocess.run(["git", "check-ignore", os.path.join("data", "agent_runs", f)],
                           cwd=repo, capture_output=True, text=True)
        assert r.returncode != 0, f"{f} is gitignored and would not deploy"
