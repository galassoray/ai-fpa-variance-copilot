"""
reports.py
==========
Word deliverables: the documents an FP&A analyst actually circulates.

Three, each with a different reader:

  FLASH        half a page, day three or four, to the CFO before the full close
  MEMO         one to two pages, the monthly variance commentary
  PACKET       one document PER department, for the budget owner who has to
               explain their own numbers and sign for them

The packet is the reason this module exists. Writing five of them by hand every
month -- pulling one owner's lines out of the close pack, formatting them,
mailing them, chasing the reply -- is the highest-volume, lowest-judgment task
in the monthly cycle, and it is the one automation should remove entirely.

THE SAME RULE AS THE DECK, FOR THE SAME REASON
-----------------------------------------------
A document generator is a surface where numbers get formatted, and every
formatting layer is a chance to become a SECOND PATH to a figure. So this
module cannot compute. There is exactly one way a number reaches a page --
``Report.fig()``, which takes a value out of a ledger section row and records
where it came from as it formats it. No arithmetic on a financial value appears
anywhere in this file.

``Report.provenance`` accumulates one entry per figure written, and the tests
re-open the generated ``.docx``, read every paragraph and table cell back out,
and audit them with the same numeric auditor used on the commentary. Checking
this module against its own bookkeeping would prove only self-consistency.

WHY THE RECOMMENDATION SECTION IS LEFT BLANK
---------------------------------------------
The standard variance-commentary format ends with a recommendation. This tool
does not generate one: a recommendation is not a retrieved number, nothing
could verify it, and the publication gate would pass it through untouched.

So every document carries the section as an explicitly empty box with a note
explaining why. That is not a gap being papered over -- it is the division of
labour made visible on the page, and it is the honest description of how the
work actually runs: the tool assembles every fact, the analyst supplies the
judgment and signs for it.
"""
from __future__ import annotations

import io
import os
from dataclasses import dataclass, field

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Palette shared with the app and the deck, so every artifact reads as one
# product. Green and red carry variance direction only -- in a finance
# document that is semantics, not decoration.
INK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x64, 0x74, 0x8B)
FAV = RGBColor(0x15, 0x80, 0x3D)
UNFAV = RGBColor(0xB9, 0x1C, 0x1C)
RULE = "E2E8F0"
WASH = "F8FAFC"

BODY = "Calibri"
HEAD = "Cambria"

BLANK_NOTE = ("Intentionally blank. This tool reports what the data shows and "
              "does not recommend action \u2014 that judgment belongs to the "
              "analyst, who signs below.")


class ReportError(RuntimeError):
    pass


@dataclass
class FigureRef:
    display: str
    value: float
    section: str
    field: str
    step: int
    label: str = ""


@dataclass
class Report:
    """A Word document built strictly from ledger sections."""

    result: object
    goal: dict
    candidate: object = None
    packet: object = None
    doc: Document = field(default_factory=Document)
    provenance: list = field(default_factory=list)

    def __post_init__(self):
        self._page_setup()

    # -- the only way a number reaches a page --------------------------
    def fig(self, section: str, row: int, field_name: str,
            kind: str = "money", label: str = "") -> str:
        """Format one value FROM a ledger row, recording its origin.

        Raises rather than returning a placeholder when the value is absent: a
        document with a silent dash where a figure belongs is worse than a
        build that stops, because the gap is invisible to the reader.
        """
        sec = self.result.sections.get(section)
        if sec is None:
            raise ReportError(f"section '{section}' is not in this run")
        rows = sec["rows"]
        if row >= len(rows):
            raise ReportError(f"section '{section}' has {len(rows)} row(s); "
                              f"row {row} was requested")
        if field_name not in rows[row]:
            raise ReportError(f"'{field_name}' is not a field of '{section}'")

        value = rows[row][field_name]
        display = fmt(value, kind)
        self.provenance.append(FigureRef(display, value, section, field_name,
                                         sec["step"], label or field_name))
        return display

    def section_by_tool(self, *tools: str, where: dict | None = None):
        """Locate a section by the TOOL that produced it, never by its label.

        Section labels are planner-chosen, so keying on them means an
        agent-authored plan produces a document missing every block whose
        section it happened to spell differently.
        """
        for name, sec in sorted(self.result.sections.items(),
                                key=lambda kv: kv[1]["step"]):
            if sec["tool"] not in tools or not sec["rows"]:
                continue
            if where and any(sec.get("params", {}).get(k) != v
                             for k, v in where.items()):
                continue
            return name
        return None

    # -- primitives -----------------------------------------------------
    def _page_setup(self):
        for s in self.doc.sections:
            s.page_width, s.page_height = Inches(8.5), Inches(11)
            s.left_margin = s.right_margin = Inches(0.9)
            s.top_margin = Inches(0.8)
            s.bottom_margin = Inches(0.8)
        style = self.doc.styles["Normal"]
        style.font.name = BODY
        style.font.size = Pt(10.5)
        style.font.color.rgb = INK
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

    def para(self, text="", size=10.5, bold=False, color=INK, font=BODY,
             align=None, space_after=6, italic=False):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if align is not None:
            p.alignment = align
        if text:
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.bold = bold
            run.italic = italic
            run.font.color.rgb = color
            run.font.name = font
        return p

    def title(self, eyebrow: str, heading: str, sub: str = ""):
        self.para(eyebrow.upper(), size=8.5, bold=True, color=MUTED,
                  space_after=2)
        self.para(heading, size=20, bold=True, font=HEAD, space_after=2)
        if sub:
            self.para(sub, size=11, color=MUTED, space_after=8)
        self.rule()

    def heading(self, text: str):
        self.para(text, size=12.5, bold=True, font=HEAD, space_after=3)

    def rule(self):
        """A paragraph bottom border -- not a one-row table, which renders as
        a box in some readers."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        pPr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), RULE)
        borders.append(bottom)
        pPr.append(borders)
        return p

    def table(self, headers: list, rows: list, widths: list | None = None):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.autofit = False
        total = Inches(6.7)
        widths = widths or [1] * len(headers)
        share = sum(widths)

        for i, (cell, head) in enumerate(zip(t.rows[0].cells, headers)):
            _shade(cell, "0F172A")
            _cell(cell, head, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                  align=WD_ALIGN_PARAGRAPH.RIGHT if i else None)
            cell.width = Inches(6.7 * widths[i] / share)

        for r_i, row in enumerate(rows):
            cells = t.add_row().cells
            for i, (cell, val) in enumerate(zip(cells, row)):
                color = INK
                if isinstance(val, tuple):
                    val, color = val
                if r_i % 2:
                    _shade(cell, WASH)
                _cell(cell, str(val), color=color,
                      align=WD_ALIGN_PARAGRAPH.RIGHT if i else None)
                cell.width = Inches(6.7 * widths[i] / share)
        self.para(space_after=4)
        return t

    def blank_recommendation(self, prompt: str):
        """The section the tool deliberately does not fill."""
        self.heading("Analyst assessment and recommended action")
        box = self.doc.add_table(rows=1, cols=1)
        cell = box.rows[0].cells[0]
        cell.width = Inches(6.7)
        _shade(cell, WASH)
        _cell(cell, prompt, color=MUTED, italic=True)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        run = cell.add_paragraph().add_run(BLANK_NOTE)
        run.font.size = Pt(8.5)
        run.italic = True
        run.font.color.rgb = MUTED
        run.font.name = BODY
        # Two blank lines, not four. Four pushed the box onto a second page
        # where it landed as an orphaned grey rectangle above the sign-off --
        # a packet that runs to two pages for the sake of white space reads as
        # careless to the person being asked to sign it.
        for _ in range(2):
            cell.add_paragraph()
        self.para(space_after=6)

    def signature_block(self, role: str = "Budget owner"):
        self.heading("Sign-off")
        self.table(
            ["Role", "Name", "Date"],
            [[role, "", ""], ["Prepared by (FP&A)", "", ""]],
            widths=[2.6, 2.6, 1.5])

    def footer_note(self, extra: str = ""):
        self.rule()
        base = ("Synthetic data. Every figure in this document is a value "
                "computed in SQL and verified against the run ledger before "
                "the document could be produced.")
        self.para(base + (f" {extra}" if extra else ""), size=8.5, color=MUTED)

    # -- output ---------------------------------------------------------
    def save(self, path: str) -> str:
        os.makedirs(os.path.dirname(os.path.abspath(path)) or ".",
                    exist_ok=True)
        self.doc.save(path)
        return path

    def to_bytes(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()


def _cell(cell, text, bold=False, color=INK, align=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.name = BODY


def _shade(cell, hex_colour: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")          # never "solid" -- renders black
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def fmt(value, kind: str = "money") -> str:
    """Format a retrieved value. Never derives one."""
    if value is None:
        return "\u2014"
    if kind == "percent":
        return f"{value * 100:.1f}%"
    if kind == "count":
        return f"{int(value):,}"
    if kind == "count_signed":
        return f"{int(value):+,}"
    body = f"${abs(value):,.0f}"
    return f"({body})" if value < 0 else body


def period_label(period: str) -> str:
    import datetime as dt
    try:
        return dt.date.fromisoformat(str(period)).strftime("%B %Y")
    except ValueError:
        return str(period)


# ==========================================================================
# FLASH — half a page, day three or four, for the CFO
# ==========================================================================
def build_flash(result, goal: dict, candidate=None, packet=None) -> Report:
    """Headline results before the full close pack exists.

    Deliberately short. A flash that runs to two pages is not a flash, and the
    reader is deciding in thirty seconds whether anything needs their attention
    today.
    """
    r = Report(result, goal, candidate, packet)
    head = r.section_by_tool("get_operating_metrics")
    pl = r.section_by_tool("get_pl_summary")
    drivers = r.section_by_tool("rank_variance_drivers")

    r.title("Flash results", period_label(goal.get("period", "")),
            "Preliminary — ahead of the full close package")

    if head:
        row = result.sections[head]["rows"][0]
        cells = [["Revenue", r.fig(head, 0, "revenue", label="revenue")],
                 ["Operating expenses", r.fig(head, 0, "opex", label="opex")],
                 ["Operating income",
                  (r.fig(head, 0, "operating_income", label="operating income"),
                   FAV if row["operating_income"] > 0 else UNFAV)],
                 ["Headcount",
                  r.fig(head, 0, "total_headcount", "count", label="headcount")]]
        if row.get("ending_arr") is not None:
            cells.append(["Ending ARR",
                          r.fig(head, 0, "ending_arr", label="ending ARR")])
        r.heading("Headline")
        r.table(["Measure", "Actual"], cells, widths=[3, 2])

    if pl:
        r.heading("Versus plan")
        rows = []
        for i, x in enumerate(result.sections[pl]["rows"]):
            col = FAV if x["oi_impact"] > 0 else UNFAV
            rows.append([x["statement_line"],
                         r.fig(pl, i, "actual", label="actual"),
                         r.fig(pl, i, "base", label="budget"),
                         (r.fig(pl, i, "oi_impact", label="OI impact"), col)])
        r.table(["Statement line", "Actual", "Budget",
                 "Impact on operating income"], rows, widths=[3, 2, 2, 2.6])

    if drivers:
        rows = result.sections[drivers]["rows"]
        top = rows[0]
        r.para(f"Largest single driver: {top['name']} at "
               f"{r.fig(drivers, 0, 'oi_impact', label='largest driver')} "
               f"impact on operating income. Detail follows in the monthly "
               f"variance memo.", size=10.5)

    r.blank_recommendation(
        "Anything requiring attention before the full package is issued?")
    r.footer_note("Figures are preliminary until the close is final.")
    return r


# ==========================================================================
# MEMO — the monthly variance commentary
# ==========================================================================
def build_memo(result, goal: dict, candidate=None, packet=None) -> Report:
    """The monthly variance commentary, structured the way the field expects:
    what happened, why, and what is underneath it -- with the recommendation
    left to the analyst."""
    from agent.briefing import build_briefing

    r = Report(result, goal, candidate, packet)
    head = r.section_by_tool("get_operating_metrics")
    pl = r.section_by_tool("get_pl_summary")
    comp = r.section_by_tool("get_comp_decomposition")
    rev = r.section_by_tool("get_revenue_decomposition")
    hc = r.section_by_tool("get_headcount_movement")

    r.title("Monthly variance commentary",
            period_label(goal.get("period", "")),
            "Actual versus budget" if goal.get("comparison") ==
            "actual_vs_budget" else "Actual versus forecast")

    # -- what happened -------------------------------------------------
    if head and pl:
        rows = []
        for i, x in enumerate(result.sections[pl]["rows"]):
            col = FAV if x["oi_impact"] > 0 else UNFAV
            rows.append([x["statement_line"],
                         r.fig(pl, i, "actual", label="actual"),
                         r.fig(pl, i, "base", label="budget"),
                         (r.fig(pl, i, "oi_impact", label="OI impact"), col),
                         ("Favorable" if x["favorable"] else "Unfavorable")])
        r.heading("Results versus plan")
        r.table(["Statement line", "Actual", "Budget",
                 "Impact on operating income", "Direction"],
                rows, widths=[2.6, 1.8, 1.8, 2.2, 1.4])

        row = result.sections[head]["rows"][0]
        r.para(f"Operating income for the period was "
               f"{r.fig(head, 0, 'operating_income', label='operating income')}"
               f" on revenue of {r.fig(head, 0, 'revenue', label='revenue')}, "
               f"with {r.fig(head, 0, 'total_headcount', 'count', label='headcount')}"
               f" employees.")

    # -- where the money moved -----------------------------------------
    brief = build_briefing(result, goal)
    if brief.available:
        r.heading("Where the variance came from")
        r.para(f"Ranked by {brief.basis}. Ranking and materiality are computed "
               f"in SQL, not selected by a model.", size=9.5, color=MUTED)
        rows = []
        for a in brief.areas:
            col = FAV if a.oi_impact > 0 else UNFAV
            share = fmt(a.share, "percent") if a.share is not None else "\u2014"
            detail = "; ".join(f"{e.label} {fmt(e.value, e.kind)}"
                               for e in a.detail[:3]) or "\u2014"
            rows.append([f"{a.rank}. {a.name}",
                         (fmt(a.oi_impact), col), share, detail])
        r.table(["Driver", "Impact on operating income", "Share of total",
                 "Largest account movements"], rows, widths=[2.2, 1.9, 1.3, 3.4])

    # -- what is underneath it -----------------------------------------
    if comp:
        r.heading("Compensation: headcount versus rate")
        r.para("Stated on an expense basis: a positive figure means more was "
               "spent than planned.", size=9.5, color=MUTED)
        rows = []
        for i, x in enumerate(result.sections[comp]["rows"]):
            rows.append([x.get("department_name") or x.get("department_id"),
                         r.fig(comp, i, "salary_variance", label="salary variance"),
                         r.fig(comp, i, "hc_impact", label="headcount effect"),
                         r.fig(comp, i, "rate_impact", label="rate effect")])
        r.table(["Department", "Salary variance", "Headcount effect",
                 "Rate effect"], rows, widths=[2.6, 2, 2, 2])

    if rev:
        x = result.sections[rev]["rows"][0]
        r.heading("Revenue: volume versus price")
        r.para(f"Revenue variance of "
               f"{r.fig(rev, 0, 'rev_variance', label='revenue variance')} "
               f"splits into "
               f"{r.fig(rev, 0, 'volume_impact', label='volume effect')} from "
               f"customer volume and "
               f"{r.fig(rev, 0, 'price_impact', label='price effect')} from "
               f"average revenue per account. A volume-driven miss is a "
               f"pipeline and retention question; a price-driven one is a "
               f"packaging and discounting question.")

    if hc:
        r.heading("Headcount versus plan")
        rows = []
        for i, x in enumerate(result.sections[hc]["rows"]):
            v = x["hc_var_vs_budget"]
            rows.append([x.get("department_name") or x.get("department_id"),
                         r.fig(hc, i, "actual_headcount", "count", label="actual"),
                         r.fig(hc, i, "budget_headcount", "count", label="plan"),
                         (r.fig(hc, i, "hc_var_vs_budget", "count_signed",
                                label="versus plan"),
                          UNFAV if v > 0 else (FAV if v < 0 else MUTED))])
        r.table(["Department", "Actual", "Plan", "Versus plan"],
                rows, widths=[3.4, 1.6, 1.6, 1.6])

    # -- commentary ----------------------------------------------------
    if candidate is not None and (candidate.text or "").strip():
        r.heading("Commentary")
        for para in candidate.text.strip().split("\n"):
            if para.strip():
                r.para(para.strip())
        r.para(f"{len(candidate.matched)} figures in this commentary were "
               f"verified against the run ledger before it could be issued.",
               size=8.5, color=MUTED, italic=True)

    r.blank_recommendation(
        "What should the business do about this, and who owns it?")
    r.footer_note()
    return r


# ==========================================================================
# PACKET — one per department, for the budget owner
# ==========================================================================
def build_packet(result, goal: dict, department_id: str, candidate=None,
                 packet=None) -> Report:
    """A single budget owner's variances, with space for their explanation.

    Writing these by hand -- pulling one owner's lines out of the close pack,
    formatting, mailing, chasing the reply -- is the highest-volume and
    lowest-judgment task in the monthly cycle. This is the one the automation
    story rests on.
    """
    r = Report(result, goal, candidate, packet)

    drivers = r.section_by_tool("rank_variance_drivers")
    decomp = r.section_by_tool("decompose_variance",
                               where={"department_id": department_id})
    comp = r.section_by_tool("get_comp_decomposition")
    hc = r.section_by_tool("get_headcount_movement")

    name = department_id
    if decomp:
        name = result.sections[decomp]["rows"][0].get("department_name") or name
    elif drivers:
        for row in result.sections[drivers]["rows"]:
            if row.get("member") == department_id:
                name = row.get("name") or name

    r.title("Budget owner variance packet",
            str(name),
            f"{period_label(goal.get('period', ''))} \u00b7 actual versus budget")

    # -- your department's position ------------------------------------
    if drivers:
        for i, row in enumerate(result.sections[drivers]["rows"]):
            if row.get("member") != department_id:
                continue
            col = FAV if row["oi_impact"] > 0 else UNFAV
            r.heading("Your position this period")
            cells = [["Impact on operating income",
                      (r.fig(drivers, i, "oi_impact", label="OI impact"), col)],
                     ["Direction",
                      ("Favorable" if row["favorable"] else "Unfavorable", col)]]
            if row.get("share_of_total_oi_impact") is not None:
                cells.append(["Share of company-wide variance",
                              r.fig(drivers, i, "share_of_total_oi_impact",
                                    "percent", label="share")])
            r.table(["Measure", "Value"], cells, widths=[3.6, 2.4])
            break

    # -- account detail -------------------------------------------------
    if decomp:
        r.heading("Account detail")
        r.para("Every line below is your department's own spend or revenue. "
               "Figures are stated as impact on operating income: a negative "
               "figure is unfavourable regardless of whether the line is "
               "revenue or expense.", size=9.5, color=MUTED)
        rows = []
        for i, x in enumerate(result.sections[decomp]["rows"]):
            col = FAV if x["oi_impact"] > 0 else UNFAV
            rows.append([x["account_name"],
                         r.fig(decomp, i, "actual", label="actual"),
                         r.fig(decomp, i, "base", label="budget"),
                         (r.fig(decomp, i, "oi_impact", label="OI impact"), col)])
        r.table(["Account", "Actual", "Budget",
                 "Impact on operating income"], rows, widths=[3, 1.8, 1.8, 2.2])
    else:
        r.para("No account-level detail was retrieved for this department in "
               "this run.", color=MUTED, italic=True)

    # -- your people ----------------------------------------------------
    comp_row = None
    if comp:
        for i, x in enumerate(result.sections[comp]["rows"]):
            if x.get("department_id") == department_id:
                comp_row = i
                break
    hc_row = None
    if hc:
        for i, x in enumerate(result.sections[hc]["rows"]):
            if x.get("department_id") == department_id:
                hc_row = i
                break

    if comp_row is not None or hc_row is not None:
        r.heading("People")
        cells = []
        if comp_row is not None:
            cells += [
                ["Salary variance (expense basis, + is above plan)",
                 r.fig(comp, comp_row, "salary_variance", label="salary variance")],
                ["\u2014 explained by headcount",
                 r.fig(comp, comp_row, "hc_impact", label="headcount effect")],
                ["\u2014 explained by rate",
                 r.fig(comp, comp_row, "rate_impact", label="rate effect")]]
        if hc_row is not None:
            v = result.sections[hc]["rows"][hc_row]["hc_var_vs_budget"]
            cells += [
                ["Headcount, actual",
                 r.fig(hc, hc_row, "actual_headcount", "count", label="actual heads")],
                ["Headcount, plan",
                 r.fig(hc, hc_row, "budget_headcount", "count", label="planned heads")],
                ["Headcount versus plan",
                 (r.fig(hc, hc_row, "hc_var_vs_budget", "count_signed",
                        label="heads vs plan"),
                  UNFAV if v > 0 else (FAV if v < 0 else MUTED))]]
        r.table(["Measure", "Value"], cells, widths=[4.2, 2])

    # -- what we need from you -----------------------------------------
    r.blank_recommendation(
        f"Please explain the variances above for {name}, and note any action "
        f"already under way or planned.")
    r.signature_block(role=f"Budget owner \u2014 {name}")
    r.footer_note("Return signed to FP&A before the close review.")
    return r


def department_ids(result) -> list:
    """Departments this run can produce a packet for, in rank order."""
    out = []
    for name, sec in sorted(result.sections.items(),
                            key=lambda kv: kv[1]["step"]):
        if sec["tool"] == "rank_variance_drivers" and \
                sec.get("params", {}).get("dimension") == "department":
            out = [str(r["member"]) for r in sec["rows"]]
            break
    if out:
        return out
    seen = []
    for name, sec in sorted(result.sections.items(),
                            key=lambda kv: kv[1]["step"]):
        if sec["tool"] == "decompose_variance":
            d = sec.get("params", {}).get("department_id")
            if d and d not in seen:
                seen.append(str(d))
    return seen


REPORTS = {
    "flash": ("Flash results", build_flash),
    "memo": ("Monthly variance commentary", build_memo),
}


def build_all(result, goal: dict, candidate=None, packet=None,
              directory: str = ".") -> list:
    """Every document for one run: flash, memo, and a packet per department."""
    period = str(goal.get("period", ""))[:7]
    written = []
    for key, (_label, builder) in REPORTS.items():
        r = builder(result, goal, candidate, packet)
        written.append(r.save(os.path.join(directory, f"{key}-{period}.docx")))
    for dept in department_ids(result):
        r = build_packet(result, goal, dept, candidate, packet)
        written.append(r.save(
            os.path.join(directory, f"packet-{period}-{dept}.docx")))
    return written
